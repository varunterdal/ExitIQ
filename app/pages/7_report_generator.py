import streamlit as st
from groq import Groq
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import io
from datetime import datetime
from utils.data_loader import load_all

st.set_page_config(page_title="Executive Report Generator",
                   page_icon="📄", layout="wide")
from utils.styles import load_css
st.markdown(load_css(), unsafe_allow_html=True)
from utils.styles import load_css
st.markdown(load_css(), unsafe_allow_html=True)
# Load data
data      = load_all()
insights  = data['insights']
themes    = data['themes']
risk      = data['risk']
attrition = data['attrition']

# Header
st.markdown(
    "<h1>📄 Executive Report Generator</h1>"
    "<p style='color:#a0aec0;margin-top:-16px;'>"
    "Generate a professional DOCX report for leadership</p>",
    unsafe_allow_html=True
)
st.divider()

# API Key input
api_key = st.text_input(
    "Enter your Groq API Key",
    type="password",
    placeholder="Paste your Groq API key here..."
)

if not api_key:
    st.info("👆 Enter your Groq API key to generate AI narrative")
    st.stop()

client = Groq(api_key=api_key)

st.success("✅ AI Connected")
st.divider()

# Report settings
st.subheader("⚙️ Report Settings")
col1, col2 = st.columns(2)

with col1:
    company_name = st.text_input("Company Name",
                                  value="Acme Corporation")
    report_title = st.text_input("Report Title",
                                  value="Employee Exit Intelligence Report")
with col2:
    prepared_by  = st.text_input("Prepared By",
                                  value="HR Analytics Team")
    report_date  = st.date_input("Report Date",
                                  value=datetime.today())

st.divider()

# Preview
st.subheader("📊 Report Preview")
col1, col2, col3, col4 = st.columns(4)
with col1:
    st.metric("Reviews Analyzed",
              f"{insights['total_reviews_analyzed']:,}")
with col2:
    st.metric("Attrition Rate",
              f"{insights['attrition_rate']}%")
with col3:
    st.metric("Positive Sentiment",
              f"{insights['overall_sentiment']['positive_pct']}%")
with col4:
    st.metric("Avg Rating",
              f"{insights['avg_overall_rating']}/5")

st.divider()

if st.button("🚀 Generate Executive Report", use_container_width=True):

    with st.spinner("🤖 Generating AI narrative and building report..."):

        context = f"""
        You are an expert HR consultant writing an executive report.
        Data:
        - Reviews analyzed: {insights['total_reviews_analyzed']}
        - Attrition rate: {insights['attrition_rate']}%
        - Avg rating: {insights['avg_overall_rating']}/5
        - Positive sentiment: {insights['overall_sentiment']['positive_pct']}%
        - Negative sentiment: {insights['overall_sentiment']['negative_pct']}%
        - Top themes: {json.dumps(insights['top_themes'])}
        - Top drivers: {json.dumps(insights['top_attrition_drivers'])}
        - Risk distribution: {json.dumps(insights['risk_distribution'])}
        - Dept attrition: {json.dumps(insights['dept_attrition'])}
        """

        try:
            exec_summary = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": context},
                    {"role": "user", "content":
                     "Write a 3-paragraph executive summary for "
                     "C-suite leadership. Be professional and concise."}
                ]
            ).choices[0].message.content

            recommendations = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": context},
                    {"role": "user", "content":
                     "Write 5 specific, actionable HR recommendations "
                     "numbered 1-5. Each recommendation should be "
                     "2-3 sentences."}
                ]
            ).choices[0].message.content

            risk_narrative = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": context},
                    {"role": "user", "content":
                     "Write a risk assessment paragraph explaining "
                     "the current attrition risk level and what it "
                     "means for the business."}
                ]
            ).choices[0].message.content

        except Exception as e:
            st.error(f"❌ AI Error: {str(e)}")
            st.stop()

        # Build DOCX
        doc = Document()

        section = doc.sections[0]
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

        # Title page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(report_title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        doc.add_paragraph()

        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run(company_name)
        company_run.bold = True
        company_run.font.size = Pt(16)

        doc.add_paragraph()

        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.add_run(
            f"Prepared by: {prepared_by}\n"
            f"Date: {report_date.strftime('%B %d, %Y')}\n"
            f"Powered by ExitIQ — AI Workforce Intelligence"
        )

        doc.add_page_break()

        # Section 1
        h1 = doc.add_heading('1. Executive Summary', level=1)
        h1.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        doc.add_paragraph(exec_summary)
        doc.add_paragraph()

        # Section 2
        h2 = doc.add_heading('2. Key Workforce Metrics', level=1)
        h2.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        metrics_table = doc.add_table(rows=5, cols=2)
        metrics_table.style = 'Table Grid'
        metrics_data = [
            ('Total Reviews Analyzed',
             str(insights['total_reviews_analyzed'])),
            ('Overall Attrition Rate',
             f"{insights['attrition_rate']}%"),
            ('Average Employee Rating',
             f"{insights['avg_overall_rating']} / 5"),
            ('Positive Sentiment',
             f"{insights['overall_sentiment']['positive_pct']}%"),
            ('Negative Sentiment',
             f"{insights['overall_sentiment']['negative_pct']}%"),
        ]
        for i, (label, value) in enumerate(metrics_data):
            metrics_table.rows[i].cells[0].text = label
            metrics_table.rows[i].cells[1].text = value
            metrics_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Section 3
        h3 = doc.add_heading('3. Top Employee Themes', level=1)
        h3.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        theme_table = doc.add_table(rows=len(themes)+1, cols=3)
        theme_table.style = 'Table Grid'
        hdr = theme_table.rows[0].cells
        hdr[0].text = 'Theme'
        hdr[1].text = 'Mentions'
        hdr[2].text = 'Avg Rating'
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
        for i, row in themes.iterrows():
            r = theme_table.rows[i+1].cells
            r[0].text = str(row['theme'])
            r[1].text = str(row['count'])
            r[2].text = str(row['avg_rating'])

        doc.add_paragraph()

        # Section 4
        h4 = doc.add_heading('4. Attrition Drivers', level=1)
        h4.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        driver_table = doc.add_table(rows=len(attrition)+1, cols=2)
        driver_table.style = 'Table Grid'
        hdr2 = driver_table.rows[0].cells
        hdr2[0].text = 'Factor'
        hdr2[1].text = 'Importance Score'
        for cell in hdr2:
            cell.paragraphs[0].runs[0].bold = True
        for i, row in attrition.iterrows():
            r = driver_table.rows[i+1].cells
            r[0].text = str(row['feature'])
            r[1].text = str(row['importance'])

        doc.add_paragraph()

        # Section 5
        h5 = doc.add_heading('5. Risk Assessment', level=1)
        h5.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        doc.add_paragraph(risk_narrative)
        doc.add_paragraph()

        risk_counts = risk['risk_level'].value_counts().reset_index()
        risk_counts.columns = ['risk_level', 'count']
        risk_table = doc.add_table(rows=len(risk_counts)+1, cols=2)
        risk_table.style = 'Table Grid'
        hdr3 = risk_table.rows[0].cells
        hdr3[0].text = 'Risk Level'
        hdr3[1].text = 'Employee Count'
        for cell in hdr3:
            cell.paragraphs[0].runs[0].bold = True
        for i, row in risk_counts.iterrows():
            r = risk_table.rows[i+1].cells
            r[0].text = str(row['risk_level'])
            r[1].text = str(row['count'])

        doc.add_paragraph()

        # Section 6
        h6 = doc.add_heading('6. Strategic Recommendations', level=1)
        h6.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        doc.add_paragraph(recommendations)
        doc.add_paragraph()

        # Footer
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"Generated by ExitIQ — AI Workforce Intelligence Platform\n"
            f"© {datetime.today().year} {company_name}. Confidential."
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Save and download
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Report generated successfully!")
        st.divider()

        st.download_button(
            label="📥 Download Executive Report (DOCX)",
            data=buffer,
            file_name=f"ExitIQ_Report_"
                      f"{datetime.today().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument"
                  ".wordprocessingml.document",
            use_container_width=True
        )

        st.balloons()